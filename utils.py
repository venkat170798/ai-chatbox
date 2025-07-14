# utils.py

import fitz  # PyMuPDF
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document

# ✅ PDF Loader & Splitter
def load_and_chunk_pdfs(file_path: str):
    documents = []
    if not file_path.lower().endswith(".pdf"):
        return documents

    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        chunks = text_splitter.split_text(text)

        for chunk in chunks:
            documents.append(Document(page_content=chunk))

        print("📄 PDF Chunks:", len(documents))
    except Exception as e:
        print(f"[PDF Error] {e}")

    return documents

# ✅ Word Loader & Splitter
def load_and_chunk_word_docs(file_path: str):
    documents = []
    if not file_path.lower().endswith(".docx"):
        return documents

    try:
        doc = docx.Document(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        chunks = text_splitter.split_text(full_text)

        for chunk in chunks:
            documents.append(Document(page_content=chunk))

        print("📄 Word Chunks:", len(documents))
    except Exception as e:
        print(f"[Word Doc Error] {e}")

    return documents
